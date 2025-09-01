#!/usr/bin/env python3
"""
Create a sample DOCX file for testing the Digital Twin Parser.
"""

from docx import Document
import os

def create_sample_docx(output_path: str):
    """Create a sample DOCX file with various content types."""
    
    # Create a new Document
    doc = Document()
    
    # Add a title
    title = doc.add_heading('MYRRHA Cryoplant Technical Requirements', 0)
    
    # Add introduction heading
    doc.add_heading('Introduction', level=1)
    
    # Add a paragraph
    intro_paragraph = doc.add_paragraph()
    intro_paragraph.add_run('This document outlines the technical requirements for the MYRRHA Cryoplant (QPLANT) based on the Deep Research report.')
    
    # Add a subheading
    doc.add_heading('System Overview', level=2)
    
    # Add another paragraph
    overview_paragraph = doc.add_paragraph()
    overview_paragraph.add_run('The cryoplant system consists of multiple components including cryomodules, cold valve boxes, and cryogenic lines.')
    
    # Add a table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Weight (kg)'
    hdr_cells[2].text = 'Temperature (K)'
    
    # Add table rows
    components = [
        ('Cryomodule (QM)', '1000', '4.5'),
        ('Cold Valve Box (QVB)', '200', '4.5'),
        ('Cryogenic lines', '50', '2.0')
    ]
    
    for component, weight, temp in components:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = weight
        row_cells[2].text = temp
    
    # Add a list
    doc.add_heading('Key Features', level=2)
    doc.add_paragraph('High efficiency cooling system', style='List Bullet')
    doc.add_paragraph('Modular design for scalability', style='List Bullet')
    doc.add_paragraph('Advanced monitoring and control', style='List Bullet')
    
    # Save the document
    doc.save(output_path)
    print(f"Sample DOCX created: {output_path}")

if __name__ == "__main__":
    os.makedirs('input', exist_ok=True)
    create_sample_docx('input/0604_parseme.docx')